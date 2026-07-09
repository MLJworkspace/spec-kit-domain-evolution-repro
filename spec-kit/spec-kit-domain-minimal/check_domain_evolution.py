#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple


VALID_ENTITY_DECISIONS = {"reused", "reused_with_modification", "new"}
VALID_ATTRIBUTE_DECISIONS = {"reused", "reused_with_modification", "new"}
VALID_RELATIONSHIP_DECISIONS = {
    "reused",
    "reused_with_modification",
    "new",
    "changed",
}

ENTITY_SYNONYMS = {
    "client": {"customer"},
    "customer": {"client"},
    "buyer": {"customer", "client"},
    "user": {"customer", "client"},
}


@dataclass
class Finding:
    area: str
    result: str
    finding: str


@dataclass
class Issue:
    severity: str
    issue: str
    suggested_fix: str


@dataclass
class WarningRow:
    warning: str
    meaning: str


def load_json(path: Path) -> Dict[str, Any]:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise SystemExit(f"Missing file: {path}") from exc
    except json.JSONDecodeError as exc:
        raise SystemExit(f"Invalid JSON in {path}: {exc}") from exc


def normalize_text(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()


def singularize(token: str) -> str:
    if len(token) > 3 and token.endswith("ies"):
        return token[:-3] + "y"
    if len(token) > 3 and token.endswith("s"):
        return token[:-1]
    return token


def tokens(value: str) -> set[str]:
    return {singularize(tok) for tok in normalize_text(value).split() if tok}


def name_key(name: str) -> str:
    return " ".join(sorted(tokens(name)))


def entity_index(model: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    return {entity.get("name", ""): entity for entity in model.get("entities", [])}


def relationship_tuple(rel: Dict[str, Any]) -> Tuple[str, str, str, str]:
    return (
        rel.get("source_entity", ""),
        normalize_verb(rel.get("verb", "")),
        rel.get("target_entity", ""),
        rel.get("cardinality", ""),
    )


def relationship_index(model: Dict[str, Any]) -> Dict[Tuple[str, str, str, str], Dict[str, Any]]:
    return {relationship_tuple(rel): rel for rel in model.get("relationships", [])}


def normalize_verb(verb: str) -> str:
    return " ".join(verb.lower().strip().split())


def relationship_label(rel_tuple: Tuple[str, str, str, str]) -> str:
    source, verb, target, cardinality = rel_tuple
    return f"{source} --{verb} [{cardinality}]--> {target}"


def entity_decisions(rationale: Dict[str, Any]) -> List[Dict[str, Any]]:
    return rationale.get("entity_decisions", []) or []


def attribute_decisions(rationale: Dict[str, Any]) -> List[Dict[str, Any]]:
    return rationale.get("attribute_decisions", []) or []


def relationship_decisions(rationale: Dict[str, Any]) -> List[Dict[str, Any]]:
    return rationale.get("relationship_decisions", []) or []


def find_entity_decision(rationale: Dict[str, Any], generated_name: str) -> Optional[Dict[str, Any]]:
    for decision in entity_decisions(rationale):
        if decision.get("generated_name") == generated_name:
            return decision
    return None


def find_relationship_decision(
    rationale: Dict[str, Any], rel_tuple: Tuple[str, str, str, str]
) -> Optional[Dict[str, Any]]:
    source, verb, target, cardinality = rel_tuple
    for decision in relationship_decisions(rationale):
        if (
            decision.get("source_entity") == source
            and normalize_verb(decision.get("verb", "")) == verb
            and decision.get("target_entity") == target
            and decision.get("cardinality") == cardinality
        ):
            return decision
    return None


def attribute_set(entity: Dict[str, Any]) -> set[str]:
    return {attr for attr in entity.get("attributes", []) if isinstance(attr, str)}


def neighbor_signature(model: Dict[str, Any], entity_name: str) -> set[str]:
    signature = set()
    for rel in model.get("relationships", []):
        source = rel.get("source_entity", "")
        target = rel.get("target_entity", "")
        verb = normalize_verb(rel.get("verb", ""))
        if source == entity_name:
            signature.add(f"out:{verb}:{target}")
        if target == entity_name:
            signature.add(f"in:{verb}:{source}")
    return signature


def jaccard(left: set[str], right: set[str]) -> float:
    if not left and not right:
        return 0.0
    return len(left & right) / len(left | right)


def synonym_overlap(left: str, right: str) -> bool:
    left_tokens = tokens(left)
    right_tokens = tokens(right)
    for token in left_tokens:
        if ENTITY_SYNONYMS.get(token, set()) & right_tokens:
            return True
    for token in right_tokens:
        if ENTITY_SYNONYMS.get(token, set()) & left_tokens:
            return True
    return False


def alias_score(
    base_model: Dict[str, Any],
    next_model: Dict[str, Any],
    base_name: str,
    next_name: str,
) -> float:
    base_entity = entity_index(base_model).get(base_name, {})
    next_entity = entity_index(next_model).get(next_name, {})
    name_overlap = jaccard(tokens(base_name), tokens(next_name))
    if synonym_overlap(base_name, next_name):
        name_overlap = max(name_overlap, 0.85)

    attr_overlap = jaccard(attribute_set(base_entity), attribute_set(next_entity))
    neighbor_overlap = jaccard(
        neighbor_signature(base_model, base_name),
        neighbor_signature(next_model, next_name),
    )
    return max(name_overlap, attr_overlap * 0.8, neighbor_overlap * 0.7)


def best_alias_candidate(
    base_model: Dict[str, Any], next_model: Dict[str, Any], next_name: str
) -> Tuple[str, float]:
    best_name = ""
    best_score = 0.0
    for base_name in entity_index(base_model):
        score = alias_score(base_model, next_model, base_name, next_name)
        if score > best_score:
            best_name = base_name
            best_score = score
    return best_name, best_score


def validate_shapes(
    base_model: Dict[str, Any],
    next_model: Dict[str, Any],
    rationale: Dict[str, Any],
    findings: List[Finding],
    issues: List[Issue],
    warnings: List[WarningRow],
) -> None:
    for label, model in (("base model", base_model), ("next model", next_model)):
        missing = [key for key in ("entities", "relationships", "warnings") if key not in model]
        if missing:
            issues.append(
                Issue(
                    "High",
                    f"The {label} is missing required fields: {', '.join(missing)}.",
                    "Regenerate or repair the extracted domain model before comparing.",
                )
            )
        else:
            findings.append(Finding("Shape", "Pass", f"The {label} has the required top-level fields."))
        for warning in model.get("warnings", []) or []:
            warnings.append(WarningRow(f"{label} parser warning", str(warning)))

    missing_rationale = [
        key
        for key in (
            "entity_decisions",
            "attribute_decisions",
            "relationship_decisions",
            "conflicts",
            "warnings",
        )
        if key not in rationale
    ]
    if missing_rationale:
        issues.append(
            Issue(
                "High",
                f"The rationale is missing required fields: {', '.join(missing_rationale)}.",
                "Regenerate the rationale with the updated specify command.",
            )
        )
    else:
        findings.append(Finding("Shape", "Pass", "The rationale has the required decision sections."))


def validate_decision_values(
    rationale: Dict[str, Any],
    findings: List[Finding],
    issues: List[Issue],
) -> None:
    invalid = []
    for decision in entity_decisions(rationale):
        value = decision.get("decision")
        if value not in VALID_ENTITY_DECISIONS:
            invalid.append(f"entity decision {value!r} for {decision.get('generated_name')!r}")
    for decision in attribute_decisions(rationale):
        value = decision.get("decision")
        if value not in VALID_ATTRIBUTE_DECISIONS:
            invalid.append(f"attribute decision {value!r} for {decision.get('entity')}.{decision.get('attribute')}")
    for decision in relationship_decisions(rationale):
        value = decision.get("decision")
        if value not in VALID_RELATIONSHIP_DECISIONS:
            invalid.append(
                f"relationship decision {value!r} for "
                f"{decision.get('source_entity')}->{decision.get('target_entity')}"
            )

    if invalid:
        issues.append(
            Issue(
                "High",
                "Invalid rationale decision values found: " + "; ".join(invalid) + ".",
                "Use only reused, reused_with_modification, new, and changed where allowed.",
            )
        )
    else:
        findings.append(Finding("Rationale", "Pass", "All rationale decision values are valid."))


def check_entities(
    base_model: Dict[str, Any],
    next_model: Dict[str, Any],
    rationale: Dict[str, Any],
    findings: List[Finding],
    issues: List[Issue],
    warnings: List[WarningRow],
) -> None:
    base_entities = entity_index(base_model)
    next_entities = entity_index(next_model)

    for next_name in sorted(next_entities):
        decision = find_entity_decision(rationale, next_name)
        decision_value = decision.get("decision") if decision else ""

        if next_name in base_entities:
            if decision_value == "new":
                issues.append(
                    Issue(
                        "Medium",
                        f"`{next_name}` exists in both models but the rationale marks it as new.",
                        "Mark it as reused or reused_with_modification.",
                    )
                )
            elif decision_value in {"reused", "reused_with_modification"}:
                findings.append(
                    Finding("Entities", "Pass", f"`{next_name}` is present in both models and rationale marks it as {decision_value}.")
                )
            else:
                issues.append(
                    Issue(
                        "Medium",
                        f"`{next_name}` exists in both models but has no matching entity rationale decision.",
                        "Add a rationale entity_decisions entry for this reused entity.",
                    )
                )
        else:
            alias_name, score = best_alias_candidate(base_model, next_model, next_name)
            if decision_value == "new":
                if score >= 0.75:
                    issues.append(
                        Issue(
                            "Medium",
                            f"`{next_name}` is marked new but may duplicate existing `{alias_name}`.",
                            "Use the existing entity name and mark the rationale as reused_with_modification if they mean the same thing.",
                        )
                    )
                else:
                    findings.append(Finding("Entities", "Pass", f"`{next_name}` is a new entity in the second model."))
            elif decision_value == "reused_with_modification":
                existing = decision.get("existing_name", "")
                if existing not in base_entities:
                    issues.append(
                        Issue(
                            "Medium",
                            f"`{next_name}` is marked reused_with_modification from `{existing}`, but `{existing}` is not in the base model.",
                            "Point existing_name to an entity that exists in the first model.",
                        )
                    )
                else:
                    findings.append(
                        Finding("Entities", "Pass", f"`{next_name}` reuses existing `{existing}` with a name or prompt-term modification.")
                    )
            elif decision_value == "reused":
                issues.append(
                    Issue(
                        "Medium",
                        f"`{next_name}` is marked reused but does not exist by that name in the first model.",
                        "Use reused_with_modification if it maps to an existing entity under a different name, or new if it is genuinely new.",
                    )
                )
            else:
                issues.append(
                    Issue(
                        "Medium",
                        f"`{next_name}` appears in the second model but is not explained by the rationale.",
                        "Add an entity_decisions entry as reused, reused_with_modification, or new.",
                    )
                )

    for base_name in sorted(set(base_entities) - set(next_entities)):
        warnings.append(
            WarningRow(
                f"`{base_name}` absent from second model",
                "Allowed. The second model is feature-scoped, so absence is not treated as deletion.",
            )
        )


def check_attributes(
    base_model: Dict[str, Any],
    next_model: Dict[str, Any],
    rationale: Dict[str, Any],
    findings: List[Finding],
    issues: List[Issue],
) -> None:
    base_entities = entity_index(base_model)
    next_entities = entity_index(next_model)

    for next_name, next_entity in sorted(next_entities.items()):
        base_attrs = attribute_set(base_entities.get(next_name, {}))
        for attr in sorted(attribute_set(next_entity)):
            decision = next(
                (
                    item
                    for item in attribute_decisions(rationale)
                    if item.get("entity") == next_name and item.get("attribute") == attr
                ),
                None,
            )
            if attr in base_attrs:
                findings.append(Finding("Attributes", "Pass", f"`{next_name}.{attr}` is reused from the first model."))
                if decision and decision.get("decision") == "new":
                    issues.append(
                        Issue(
                            "Low",
                            f"`{next_name}.{attr}` exists in the first model but rationale marks it as new.",
                            "Mark the attribute decision as reused.",
                        )
                    )
            else:
                if decision and decision.get("decision") == "new":
                    findings.append(Finding("Attributes", "Pass", f"`{next_name}.{attr}` is a new attribute in the second model."))
                elif decision:
                    findings.append(
                        Finding("Attributes", "Warning", f"`{next_name}.{attr}` is absent from the first model but rationale marks it as {decision.get('decision')}.")
                    )
                else:
                    issues.append(
                        Issue(
                            "Low",
                            f"`{next_name}.{attr}` appears in the second model but is not explained by the rationale.",
                            "Add an attribute_decisions entry or remove the attribute if it is accidental.",
                        )
                    )


def same_endpoints(left: Tuple[str, str, str, str], right: Tuple[str, str, str, str]) -> bool:
    return left[0] == right[0] and left[2] == right[2]


def check_relationships(
    base_model: Dict[str, Any],
    next_model: Dict[str, Any],
    rationale: Dict[str, Any],
    findings: List[Finding],
    issues: List[Issue],
) -> None:
    base_rels = relationship_index(base_model)
    next_rels = relationship_index(next_model)
    base_entities = entity_index(base_model)
    next_entities = entity_index(next_model)

    for rel_tuple in sorted(next_rels):
        decision = find_relationship_decision(rationale, rel_tuple)
        label = relationship_label(rel_tuple)

        if rel_tuple in base_rels:
            findings.append(Finding("Relationships", "Pass", f"`{label}` is reused from the first model."))
            if decision and decision.get("decision") == "new":
                issues.append(
                    Issue(
                        "Medium",
                        f"`{label}` exists in the first model but rationale marks it as new.",
                        "Mark the relationship decision as reused.",
                    )
                )
            continue

        same_stv = [
            base_tuple
            for base_tuple in base_rels
            if base_tuple[0] == rel_tuple[0] and base_tuple[1] == rel_tuple[1] and base_tuple[2] == rel_tuple[2]
        ]
        if same_stv:
            base_label = relationship_label(same_stv[0])
            issues.append(
                Issue(
                    "High",
                    f"`{label}` changes cardinality of existing `{base_label}`.",
                    "Mark this as changed in the rationale and confirm the change is intentional.",
                )
            )
            continue

        same_pair = [base_tuple for base_tuple in base_rels if same_endpoints(base_tuple, rel_tuple)]
        if same_pair and rel_tuple[0] in base_entities and rel_tuple[2] in base_entities:
            findings.append(
                Finding("Relationships", "Warning", f"`{label}` connects existing entities with new wording.")
            )
            continue

        endpoint_status = []
        for entity_name in (rel_tuple[0], rel_tuple[2]):
            if entity_name in base_entities:
                endpoint_status.append("reused")
            elif entity_name in next_entities:
                endpoint_status.append("new")
            else:
                endpoint_status.append("missing")

        if "missing" in endpoint_status:
            issues.append(
                Issue(
                    "Medium",
                    f"`{label}` references an entity missing from the second model.",
                    "Ensure both relationship endpoints are included in the extracted model.",
                )
            )
        else:
            findings.append(Finding("Relationships", "Pass", f"`{label}` is an extension in the second model."))

        if not decision:
            issues.append(
                Issue(
                    "Low",
                    f"`{label}` appears in the second model but is not explained by the rationale.",
                    "Add a relationship_decisions entry for this relationship.",
                )
            )

    for decision in relationship_decisions(rationale):
        rel_tuple = (
            decision.get("source_entity", ""),
            normalize_verb(decision.get("verb", "")),
            decision.get("target_entity", ""),
            decision.get("cardinality", ""),
        )
        if rel_tuple not in next_rels:
            issues.append(
                Issue(
                    "Medium",
                    f"The rationale claims `{relationship_label(rel_tuple)}`, but the extracted second model does not contain it.",
                    "Revise the spec into a parser-readable relationship form or remove the rationale claim.",
                )
            )


def check_conflicts(
    base_model: Dict[str, Any],
    next_model: Dict[str, Any],
    rationale: Dict[str, Any],
    findings: List[Finding],
    issues: List[Issue],
) -> None:
    conflict_count = sum(1 for issue in issues if "changes cardinality" in issue.issue)
    rationale_conflicts = rationale.get("conflicts", []) or []

    if conflict_count == 0:
        findings.append(Finding("Conflicts", "Pass", "No contradictory cardinalities were found for existing relationships."))
    elif not rationale_conflicts:
        issues.append(
            Issue(
                "High",
                "A deterministic conflict was found but the rationale conflicts array is empty.",
                "Add a conflicts entry explaining whether this is explicit evolution or needs clarification.",
            )
        )
    else:
        findings.append(Finding("Conflicts", "Warning", "Conflicts were found and the rationale includes conflict entries."))


def classify_overall(findings: Sequence[Finding], issues: Sequence[Issue]) -> str:
    if any(issue.severity == "High" for issue in issues):
        return "Conflict or invalid comparison"
    if issues:
        return "Extension with issues to review"
    has_modification = any("reused_with_modification" in finding.finding for finding in findings)
    has_new = any("new entity" in finding.finding or "extension" in finding.finding for finding in findings)
    if has_modification and has_new:
        return "Extension with reused/modified concepts"
    if has_new:
        return "Clean extension"
    if has_modification:
        return "Reuse with modification"
    return "Consistent reuse"


def markdown_table(headers: Sequence[str], rows: Sequence[Sequence[str]]) -> str:
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(str(cell).replace("\n", " ") for cell in row) + " |")
    return "\n".join(lines)


def build_markdown(
    overall: str,
    findings: Sequence[Finding],
    issues: Sequence[Issue],
    warnings: Sequence[WarningRow],
) -> str:
    finding_rows = [(item.area, item.result, item.finding) for item in findings]
    issue_rows = [(item.severity, item.issue, item.suggested_fix) for item in issues]
    warning_rows = [(item.warning, item.meaning) for item in warnings]

    parts = [
        "# Domain Evolution Check Report",
        "",
        f"**Overall result:** {overall}",
        "",
        markdown_table(("Area", "Result", "Finding"), finding_rows or [("Checks", "Pass", "No findings.")]),
        "",
        "## Issues",
        "",
        markdown_table(("Severity", "Issue", "Suggested fix"), issue_rows or [("None", "No issues found.", "")]),
        "",
        "## Warnings",
        "",
        markdown_table(("Warning", "Meaning"), warning_rows or [("None", "No warnings.")]),
        "",
    ]
    return "\n".join(parts)


def clean_docx_text(value: Any) -> str:
    return str(value).replace("`", "")


def set_repeat_table_header(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell: Any, width: Any) -> None:
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        from docx.oxml import OxmlElement

        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w", str(width.twips))
    tc_w.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "dxa")


def format_cell(cell: Any, value: Any, font_size: Any, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(clean_docx_text(value))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = font_size


def add_table(
    document: Any,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[Any],
    font_size: Any,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_width(header_cells[idx], widths[idx])
        format_cell(header_cells[idx], header, font_size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            format_cell(cells[idx], value, font_size)


def write_docx(
    path: Path,
    overall: str,
    findings: Sequence[Finding],
    issues: Sequence[Issue],
    warnings: Sequence[WarningRow],
) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.shared import Pt
        from docx.shared import Inches
    except ImportError as exc:
        raise SystemExit("python-docx is required for --out-docx.") from exc

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for style_name in ("Heading 1", "Heading 2"):
        styles[style_name].font.name = "Arial"

    title = document.add_paragraph()
    title_run = title.add_run("Domain Evolution Check Report")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)

    summary = document.add_paragraph()
    summary.add_run("Overall result: ").bold = True
    summary.add_run(overall)

    document.add_heading("Checks", level=1)
    add_table(
        document,
        ("Area", "Result", "Finding"),
        [(item.area, item.result, item.finding) for item in findings] or [("Checks", "Pass", "No findings.")],
        (Inches(1.2), Inches(0.8), Inches(8.0)),
        Pt(8),
    )

    document.add_heading("Issues", level=1)
    add_table(
        document,
        ("Severity", "Issue", "Suggested fix"),
        [(item.severity, item.issue, item.suggested_fix) for item in issues]
        or [("None", "No issues found.", "")],
        (Inches(1.0), Inches(4.6), Inches(4.4)),
        Pt(8),
    )

    document.add_heading("Warnings", level=1)
    add_table(
        document,
        ("Warning", "Meaning"),
        [(item.warning, item.meaning) for item in warnings] or [("None", "No warnings.")],
        (Inches(3.2), Inches(6.8)),
        Pt(8),
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def run_check(
    base_model: Dict[str, Any],
    next_model: Dict[str, Any],
    rationale: Dict[str, Any],
) -> Tuple[str, List[Finding], List[Issue], List[WarningRow]]:
    findings: List[Finding] = []
    issues: List[Issue] = []
    warnings: List[WarningRow] = []

    validate_shapes(base_model, next_model, rationale, findings, issues, warnings)
    validate_decision_values(rationale, findings, issues)
    check_entities(base_model, next_model, rationale, findings, issues, warnings)
    check_attributes(base_model, next_model, rationale, findings, issues)
    check_relationships(base_model, next_model, rationale, findings, issues)
    check_conflicts(base_model, next_model, rationale, findings, issues)

    overall = classify_overall(findings, issues)
    return overall, findings, issues, warnings


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Check whether one extracted domain model extends or modifies another."
    )
    parser.add_argument("--base", required=True, type=Path, help="Base feature-domain.json")
    parser.add_argument("--next", required=True, type=Path, help="Next feature-domain.json")
    parser.add_argument("--rationale", required=True, type=Path, help="domain-rationale.json for the next model")
    parser.add_argument("--out-md", type=Path, help="Write Markdown report")
    parser.add_argument("--out-docx", type=Path, help="Write Word .docx report")
    parser.add_argument(
        "--fail-on-issues",
        action="store_true",
        help="Exit with status 1 when the report contains issues.",
    )
    args = parser.parse_args()

    base_model = load_json(args.base)
    next_model = load_json(args.next)
    rationale = load_json(args.rationale)

    overall, findings, issues, warnings = run_check(base_model, next_model, rationale)
    markdown = build_markdown(overall, findings, issues, warnings)

    if args.out_md:
        args.out_md.parent.mkdir(parents=True, exist_ok=True)
        args.out_md.write_text(markdown, encoding="utf-8")
    if args.out_docx:
        write_docx(args.out_docx, overall, findings, issues, warnings)

    print(markdown)
    if issues and args.fail_on_issues:
        sys.exit(1)


if __name__ == "__main__":
    main()
